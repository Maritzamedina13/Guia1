"""
Generador DOCX – Guía N°1 Visita de Inicio ITM (FDE-074 V11)
Incluye todas las opciones de listas desplegables documentadas.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores institucionales ──────────────────────────────────────────────────
AZUL      = RGBColor(0x10, 0x2D, 0x69)
AZUL_L    = RGBColor(0xE8, 0xED, 0xF5)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO     = RGBColor(0x1A, 0x1A, 0x1A)
GRIS_L    = RGBColor(0xF4, 0xF6, 0xF9)
GRIS_M    = RGBColor(0xC5, 0xCD, 0xD8)

# ── Helpers XML ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color='C5CDD8'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_no_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','bottom','left','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

def cell_para(cell, text, bold=False, italic=False, size=9, color=NEGRO,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = 'Arial'
    return p

def add_run(para, text, bold=False, italic=False, size=9, color=NEGRO):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = 'Arial'
    return run

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

# ── Barra de sección ─────────────────────────────────────────────────────────
def add_section_bar(doc, number, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, AZUL)
    set_cell_no_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(f'{number}  {title}')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = BLANCO
    r.font.name = 'Arial'
    doc.add_paragraph()  # spacer

# ── Fila etiqueta + valor ────────────────────────────────────────────────────
def add_label_row(table, row_idx, label, value='', alt=False):
    row  = table.rows[row_idx]
    lcell = row.cells[0]
    vcell = row.cells[1]
    set_cell_bg(lcell, AZUL_L)
    set_cell_bg(vcell, GRIS_L if alt else BLANCO)
    for c in [lcell, vcell]:
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
    cell_para(lcell, label, bold=True, color=AZUL, size=9)
    cell_para(vcell, value, size=9)

# ── Fila de encabezado de columna ────────────────────────────────────────────
def add_header_row(table, row_idx, *texts):
    row = table.rows[row_idx]
    for i, t in enumerate(texts):
        c = row.cells[i]
        set_cell_bg(c, AZUL)
        set_cell_borders(c, top='single', bottom='single', left='single', right='single', color='102D69')
        cell_para(c, t, bold=True, color=BLANCO, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Checkbox text ────────────────────────────────────────────────────────────
CHK_ON  = '☑'
CHK_OFF = '☐'

def check_row(table, row_idx, label, checked='off', alt=False):
    row = table.rows[row_idx]
    c0  = row.cells[0]
    c1  = row.cells[1]
    set_cell_bg(c0, GRIS_L if alt else BLANCO)
    set_cell_bg(c1, GRIS_L if alt else BLANCO)
    for c in [c0, c1]:
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
    cell_para(c0, label, size=9)
    sym = CHK_ON if checked == 'on' else CHK_OFF
    cell_para(c1, sym + '  Sí       ' + CHK_OFF + '  No       ' + CHK_OFF + '  N/A',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════════════
#  DOCUMENTO PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes de página carta
sec = doc.sections[0]
sec.page_width  = Cm(21.59)
sec.page_height = Cm(27.94)
sec.left_margin   = Cm(1.8)
sec.right_margin  = Cm(1.8)
sec.top_margin    = Cm(1.8)
sec.bottom_margin = Cm(1.5)

# Estilo de párrafo base
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(9)

# ── ENCABEZADO ───────────────────────────────────────────────────────────────
hdr_tbl = doc.add_table(rows=3, cols=3)
hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
# Fondo blanco todas las celdas
for row in hdr_tbl.rows:
    for cell in row.cells:
        set_cell_bg(cell, BLANCO)

# Logo (col 0, todas las filas)
logo_cell = hdr_tbl.cell(0, 0)
logo_cell.merge(hdr_tbl.cell(2, 0))
try:
    p = logo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('Imagenes/LOGO ITM.png', width=Cm(2.8))
except Exception:
    cell_para(logo_cell, 'ITM', bold=True, size=10, color=AZUL,
              align=WD_ALIGN_PARAGRAPH.CENTER)

# Título central
tit_cell = hdr_tbl.cell(0, 1)
tit_cell.merge(hdr_tbl.cell(1, 1))
p = tit_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
r1 = p.add_run('GUÍA N°1 – VISITA DE INICIO\n')
r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = AZUL; r1.font.name = 'Arial'
r2 = p.add_run('MODALIDAD PRÁCTICA PROFESIONAL\n')
r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = AZUL; r2.font.name = 'Arial'
r3 = p.add_run('Vicerrectoría de Docencia · Dirección de Gestión Académica · Oficina de Prácticas Profesionales ITM')
r3.bold = True; r3.font.size = Pt(8); r3.font.color.rgb = AZUL; r3.font.name = 'Arial'

sub_cell = hdr_tbl.cell(2, 1)
p = sub_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Código (col 2, todas las filas)
cod_cell = hdr_tbl.cell(0, 2)
cod_cell.merge(hdr_tbl.cell(2, 2))
p = cod_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(2)
for label, val in [('Código: ', 'FDE 074'), ('Versión: ', '11'), ('Fecha: ', '01-06-2026')]:
    r_l = p.add_run(label)
    r_l.bold = True; r_l.font.size = Pt(8.5); r_l.font.color.rgb = AZUL; r_l.font.name = 'Arial'
    r_v = p.add_run(val + '\n')
    r_v.font.size = Pt(8.5); r_v.font.color.rgb = NEGRO; r_v.font.name = 'Arial'

# Borde inferior encabezado
for row in hdr_tbl.rows:
    for cell in row.cells:
        set_cell_borders(cell, bottom='single', color='102D69')

set_col_widths(hdr_tbl, [3.0, 11.5, 3.5])

# ── OBJETIVO DEL INSTRUMENTO ─────────────────────────────────────────────────
doc.add_paragraph()
p_obj = doc.add_paragraph()
p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p_obj.add_run('Objetivo del instrumento: ')
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
r2 = p_obj.add_run(
    'Registrar y documentar la visita de inicio de la práctica profesional, verificando las condiciones del '
    'escenario, identificando al practicante y al tutor, acordando el plan de trabajo y los compromisos de '
    'las partes, en cumplimiento del Decreto 0223 del 5 de marzo de 2026 y la Política de Prácticas '
    'Profesionales del ITM.')
r2.font.size = Pt(9); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'

# ── REF. NORMATIVA ───────────────────────────────────────────────────────────
p_ref = doc.add_paragraph()
p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p_ref.add_run('Ref. normativa: ')
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
r2 = p_ref.add_run(
    'Decreto 0223/2026 (Min. Trabajo) · Ley 1780/2016 · Ley 2466/2025 (Reforma Laboral) · '
    'Ley 2365/2024 · Ley 2039/2020 · Art. 81 CST · Política de Prácticas Profesionales ITM.')
r2.font.size = Pt(9); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'


# ════════════════════════════════════════════════════════════════════════════
# ══ 2. IDENTIFICACIÓN DEL PRACTICANTE ═══════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
add_section_bar(doc, '2.', 'IDENTIFICACIÓN DEL PRACTICANTE')

pract_rows = [
    ('Nombre completo:', '', False),
    ('Número de identificación:', '', True),
    ('Correo electrónico institucional:', '', False),
    ('Teléfono de contacto:', '', True),
    ('Facultad:', '', False),
    ('Programa académico:', '', True),
    ('Semestre en curso:', '', False),
    ('Monitor asignado (ITM):', '', True),
    ('Fecha de inicio de la práctica:', '', False),
    ('Fecha estimada de finalización:', '', True),
    ('Modalidad de práctica:', '', False),
]

tbl2 = doc.add_table(rows=len(pract_rows), cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl2, [7.0, 11.0])
for i, (lbl, val, alt) in enumerate(pract_rows):
    add_label_row(tbl2, i, lbl, val, alt)
set_row_height(tbl2.rows[10], 1.0)

# Opciones desplegables – Facultad
p_fac = tbl2.rows[4].cells[1].paragraphs[0]
p_fac.clear()
r = p_fac.add_run('[Lista desplegable – Seleccione facultad]\n')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
for opt in ['Facultad de Artes y Humanidades',
            'Facultad de Ciencias Económicas y Administrativas',
            'Facultad de Ciencias Exactas y Aplicadas',
            'Facultad de Ingenierías']:
    r2 = p_fac.add_run(f'  {CHK_OFF} {opt}\n')
    r2.font.size = Pt(8); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'

# Opciones desplegables – Programa (agrupadas)
programas = {
    'Artes y Humanidades': ['Artes Visuales','Cine','Artes de la Grabación y Producción Musical',
        'Tecnología en Informática Musical','Música','Licenciatura en Música'],
    'Ciencias Económicas': ['Administración de Empresas','Contaduría Pública',
        'Tecnología en Gestión de Empresas Turísticas','Negocios Internacionales','Mercadeo'],
    'Ciencias Exactas': ['Biología','Matemáticas','Física','Tecnología en Higiene y Seguridad Industrial',
        'Ingeniería Biomédica','Ingeniería Ambiental'],
    'Ingenierías': ['Ingeniería de Sistemas','Ingeniería Electrónica','Ingeniería Eléctrica',
        'Ingeniería Industrial','Ingeniería Mecatrónica','Ingeniería Mecánica',
        'Ingeniería Financiera','Ingeniería de Producción',
        'Tecnología en Electrónica','Tecnología en Sistemas Informáticos',
        'Tecnología en Gestión Industrial']
}
p_prog = tbl2.rows[5].cells[1].paragraphs[0]
p_prog.clear()
r = p_prog.add_run('[Lista desplegable – filtrada por Facultad seleccionada]\n')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
for fac, progs in programas.items():
    r_h = p_prog.add_run(f'  ── {fac}: ')
    r_h.bold = True; r_h.font.size = Pt(8); r_h.font.color.rgb = AZUL; r_h.font.name = 'Arial'
    r_p = p_prog.add_run(', '.join(progs) + '\n')
    r_p.font.size = Pt(8); r_p.font.color.rgb = NEGRO; r_p.font.name = 'Arial'

# Opciones – Semestre
p_sem = tbl2.rows[6].cells[1].paragraphs[0]
p_sem.clear()
r = p_sem.add_run('[Lista desplegable]  ')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
semestres = [f'{i}° Semestre' for i in range(1, 11)]
r2 = p_sem.add_run(' / '.join(semestres))
r2.font.size = Pt(8); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'

# Opciones – Monitor
monitores = ['Adriana Maria Zapata Cano','Alberto Alexander Fonnegra Montoya',
    'Alicia Magnolia Vasquez Franco','Angie Paola Meneses Calderon',
    'Carmen Elisa Escobar Agudelo','Carmen Liliana Barragan Guevara',
    'Claudia Ester Loaiza Alzate','Cristian David Bocanegra Villazon',
    'Erick Uribe Lopez','Estefanía Hidalgo Vásquez','Gladys Cristina Perez Rojas',
    'Guillermo Leon Vasquez Zapata','Hector Darío Calle Cano',
    'Juan David Aguirre Ramirez','Juan Sebastian Gonzalez Montoya',
    'Julie Bibiana Gomez Galvis','Kelly Mabel Ortega Zapata',
    'Luis Santiago Ramirez Alvarez','Maria Isabel Vicente Jimenez',
    'Maritza Medina Giraldo','Monica Maria Meneses Martinez',
    'Natalia Eugenia Fonnegra Gomez','Norida Astrid Valencia Bustamante',
    'Robert NG Henao','Yanneth Del Socorro Lopez']
p_mon = tbl2.rows[7].cells[1].paragraphs[0]
p_mon.clear()
r = p_mon.add_run('[Lista desplegable – Monitores asignados ITM]\n')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
for m in monitores:
    r2 = p_mon.add_run(f'  {CHK_OFF} {m}\n')
    r2.font.size = Pt(8); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'

# Modalidad de práctica
p_mod = tbl2.rows[10].cells[1].paragraphs[0]
p_mod.clear()
r = p_mod.add_run(f'{CHK_OFF} Práctica Profesional     {CHK_OFF} Práctica Social     {CHK_OFF} Práctica Internacional')
r.font.size = Pt(9); r.font.color.rgb = NEGRO; r.font.name = 'Arial'


# ════════════════════════════════════════════════════════════════════════════
# ══ 3. IDENTIFICACIÓN DEL ESCENARIO ═════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '3.', 'IDENTIFICACIÓN DEL ESCENARIO DE PRÁCTICA / EMPRESA PATROCINADORA')

esc_rows = [
    ('Razón social / Nombre del escenario:', '', False),
    ('NIT / Documento de identificación:', '', True),
    ('Tipo de escenario:', '', False),
    ('Sector económico:', '', True),
    ('Dirección / Ciudad:', '', False),
    ('Nombre del Tutor designado:', '', True),
    ('Cargo del Tutor:', '', False),
    ('Correo electrónico del Tutor:', '', True),
    ('Teléfono del Tutor:', '', False),
    ('Área / Dependencia de práctica:', '', True),
    ('Modalidad de trabajo:', '', False),
]
tbl3 = doc.add_table(rows=len(esc_rows), cols=2)
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl3, [7.0, 11.0])
for i, (lbl, val, alt) in enumerate(esc_rows):
    add_label_row(tbl3, i, lbl, val, alt)

# Tipo de escenario
p_tipo = tbl3.rows[2].cells[1].paragraphs[0]
p_tipo.clear()
r = p_tipo.add_run('[Lista desplegable]  ')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
r2 = p_tipo.add_run(f'{CHK_OFF} Empresa patrocinadora (Contrato aprendizaje)     {CHK_OFF} Escenario de práctica laboral (Vinc. Formativa)')
r2.font.size = Pt(8); r2.font.color.rgb = NEGRO; r2.font.name = 'Arial'

# Sector económico
sectores = {
    'Tecnología e Innovación': ['TI y Software','Telecomunicaciones y Redes','Electrónica y Automatización','IA y Datos'],
    'Industria y Manufactura': ['Manufactura General','Metalmecánica','Textil y Confección','Alimentaria','Química','Farmacéutica'],
    'Construcción e Infraestructura': ['Construcción y Obras Civiles','Arquitectura','Inmobiliario'],
    'Salud y Bienestar': ['Salud (Clínicas, EPS)','Tecnología Biomédica','Deporte','Medio Ambiente'],
    'Comercio y Servicios': ['Comercio Mayor/Menor','Logística y Transporte','Turismo','Publicidad','Consultoría'],
    'Finanzas y Administración': ['Banca y Seguros','Contabilidad y Auditoría','Gestión Administrativa'],
    'Educación y Cultura': ['Educación','Arte y Diseño','Medios de Comunicación','Música y Audiovisual'],
    'Sector Público': ['Entidades Gubernamentales','Organismos de Control','ONG'],
    'Energía y Recursos': ['Energía Eléctrica y Renovables','Petróleo y Gas','Gestión del Agua'],
    'Otro': ['Otro sector']
}
p_sec = tbl3.rows[3].cells[1].paragraphs[0]
p_sec.clear()
r = p_sec.add_run('[Lista desplegable – Sectores económicos]\n')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x80,0x80,0x80); r.font.name = 'Arial'
for cat, opts in sectores.items():
    r_c = p_sec.add_run(f'  ── {cat}: ')
    r_c.bold = True; r_c.font.size = Pt(8); r_c.font.color.rgb = AZUL; r_c.font.name = 'Arial'
    r_o = p_sec.add_run(', '.join(opts) + '\n')
    r_o.font.size = Pt(8); r_o.font.color.rgb = NEGRO; r_o.font.name = 'Arial'

# Modalidad de trabajo
p_mtrab = tbl3.rows[10].cells[1].paragraphs[0]
p_mtrab.clear()
r = p_mtrab.add_run(f'{CHK_OFF} Presencial     {CHK_OFF} Híbrida     {CHK_OFF} Virtual')
r.font.size = Pt(9); r.font.color.rgb = NEGRO; r.font.name = 'Arial'


# ════════════════════════════════════════════════════════════════════════════
# ══ 4. VERIFICACIÓN DE DOCUMENTOS ═══════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '4.', 'VERIFICACIÓN DE DOCUMENTOS DE VINCULACIÓN')
p_nota = doc.add_paragraph('(Decreto 0223/2026 – Arts. 2.2.6.3.2.3, 2.2.6.3.2.4 y 2.2.6.3.3.14)')
p_nota.runs[0].italic = True; p_nota.runs[0].font.size = Pt(8); p_nota.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)

docs_req = [
    'Contrato de aprendizaje / Acuerdo de voluntades / Acto administrativo',
    'Afiliación a ARL — obligatoria (Art. 2.2.6.3.2.9 y 2.2.6.3.3.7 Decreto 0223/2026)',
    'Afiliación a EPS (Salud)',
    'Afiliación a Fondo de Pensión (solo fase práctica contrato aprendizaje)',
    'Carta de inicio emitida por la OPP-ITM',
    'Reglamento interno de trabajo entregado al aprendiz (contrato aprendizaje)',
    'Aprobación de funciones por Monitor ITM',
    'Acta de inicio suscrita por las tres partes',
    'Autorización Inspector de Trabajo (solo si el practicante tiene 15–17 años – Art. 2.2.6.3.1.4)',
]
tbl4 = doc.add_table(rows=len(docs_req)+1, cols=2)
tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl4, [13.5, 4.5])
add_header_row(tbl4, 0, 'Documento requerido', 'Verificado')
for i, d in enumerate(docs_req):
    row = tbl4.rows[i+1]
    alt = i % 2 == 1
    set_cell_bg(row.cells[0], GRIS_L if alt else BLANCO)
    set_cell_bg(row.cells[1], GRIS_L if alt else BLANCO)
    for c in row.cells:
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
    cell_para(row.cells[0], d, size=9)
    cell_para(row.cells[1], f'{CHK_OFF} Sí    {CHK_OFF} No    {CHK_OFF} N/A',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


# ════════════════════════════════════════════════════════════════════════════
# ══ 5. CONDICIONES DEL ESCENARIO ════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '5.', 'VERIFICACIÓN DE CONDICIONES DEL ESCENARIO DE PRÁCTICA')
p_nota2 = doc.add_paragraph('(Decreto 0223/2026 – Art. 2.2.6.3.1.10 | Política ITM – Cap. VI, num. 6.2)')
p_nota2.runs[0].italic = True; p_nota2.runs[0].font.size = Pt(8); p_nota2.runs[0].font.color.rgb = RGBColor(0x60,0x60,0x60)

condiciones = [
    'El practicante cuenta con espacio físico/virtual adecuado para realizar sus funciones',
    'Existe un tutor designado por el escenario que acompañará al practicante',
    'Las funciones asignadas corresponden al perfil académico del programa',
    'El horario de práctica es compatible con las actividades académicas del practicante',
    'El escenario garantiza condiciones de Seguridad y Salud en el Trabajo (SST)',
    'El practicante está afiliado al SGSSS (EPS + ARL) antes de iniciar funciones',
    'El escenario cuenta con herramientas o equipos necesarios para la práctica',
    'Se han socializado los derechos y deberes del practicante',
    'El escenario conoce la Política de Prácticas ITM y el marco normativo aplicable',
    'Las actividades NO implican riesgos para menores de edad (si aplica – Art. 2.2.6.3.1.4)',
]
tbl5 = doc.add_table(rows=len(condiciones)+1, cols=2)
tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl5, [13.0, 5.0])
add_header_row(tbl5, 0, 'Condición verificada', 'Estado')
for i, cond in enumerate(condiciones):
    row = tbl5.rows[i+1]
    alt = i % 2 == 1
    set_cell_bg(row.cells[0], GRIS_L if alt else BLANCO)
    set_cell_bg(row.cells[1], GRIS_L if alt else BLANCO)
    for c in row.cells:
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
    cell_para(row.cells[0], cond, size=9)
    cell_para(row.cells[1], f'{CHK_OFF} Cumple     {CHK_OFF} No cumple',
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


# ════════════════════════════════════════════════════════════════════════════
# ══ 6. PLAN DE TRABAJO ══════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '6.', 'PLAN DE TRABAJO INICIAL (Acuerdo tripartito)')

p_nota3 = doc.add_paragraph('Objetivos formativos acordados (Practicante · Tutor · Monitor):')
p_nota3.runs[0].bold = True; p_nota3.runs[0].font.size = Pt(9); p_nota3.runs[0].font.color.rgb = AZUL

for n in range(1, 4):
    p_obj_line = doc.add_paragraph(f'{n}. _____________________________________________________________________________________________________________')
    p_obj_line.paragraph_format.space_before = Pt(0)
    p_obj_line.paragraph_format.space_after  = Pt(0)

tbl6 = doc.add_table(rows=8, cols=4)
tbl6.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl6, [1.8, 5.8, 4.3, 6.1])
add_header_row(tbl6, 0, 'Período', 'Actividades / Funciones a desarrollar',
               'Resultados de aprendizaje esperados', 'Tipo de producto')

meses = ['Mes 1','Mes 2','Mes 3','Mes 4','Mes 5','Mes 6']
for i, mes in enumerate(meses):
    row = tbl6.rows[i+1]
    alt = i % 2 == 1
    set_cell_bg(row.cells[0], GRIS_L if alt else AZUL_L)
    set_cell_bg(row.cells[1], GRIS_L if alt else BLANCO)
    set_cell_bg(row.cells[2], GRIS_L if alt else BLANCO)
    set_cell_bg(row.cells[3], GRIS_L if alt else BLANCO)
    for c in row.cells:
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
    cell_para(row.cells[0], mes, bold=True, color=AZUL, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(row.cells[1], '', size=9)
    cell_para(row.cells[2], '', size=9)
    p_tp = row.cells[3].paragraphs[0]
    r = p_tp.add_run(f'{CHK_OFF} Entregable\n{CHK_OFF} Informe\n{CHK_OFF} Ninguno')
    r.font.size = Pt(8.5); r.font.color.rgb = NEGRO; r.font.name = 'Arial'
    set_row_height(row, 1.5)

# Mayor a 6 meses
row_ext = tbl6.rows[7]
for c in row_ext.cells:
    set_cell_bg(c, BLANCO)
    set_cell_borders(c, top='single', bottom='single', left='single', right='single')
cell_para(row_ext.cells[0], 'Mayor a\n6 meses', bold=True, color=AZUL, size=8.5,
          align=WD_ALIGN_PARAGRAPH.CENTER)
row_ext.cells[1].merge(row_ext.cells[2]).merge(row_ext.cells[3])
cell_para(row_ext.cells[1], '', size=9)
set_row_height(row_ext, 1.0)

p_obs_pt = doc.add_paragraph()
r = p_obs_pt.add_run('Observaciones Monitor / Tutor: ')
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
p_obs_pt.add_run('___________________________________________________________________________________________________________')


# ════════════════════════════════════════════════════════════════════════════
# ══ 7. HORARIO ══════════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '7.', 'HORARIO DE PRÁCTICA')

tbl7 = doc.add_table(rows=1, cols=3)
tbl7.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl7, [3.0, 5.0, 10.0])
row7 = tbl7.rows[0]
for c in row7.cells:
    set_cell_bg(c, BLANCO)
    set_cell_borders(c, top='single', bottom='single', left='single', right='single')
set_cell_bg(row7.cells[0], AZUL)
cell_para(row7.cells[0], 'TOTAL SEMANAL', bold=True, color=BLANCO, size=9)
cell_para(row7.cells[1], '______  horas por semana', size=9)
p7 = row7.cells[2].paragraphs[0]
r = p7.add_run(f'Compatible con actividades académicas ITM:  {CHK_OFF} Sí     {CHK_OFF} No')
r.font.size = Pt(9); r.font.color.rgb = NEGRO; r.font.name = 'Arial'
set_row_height(row7, 0.8)

p_hor_obs = doc.add_paragraph()
r = p_hor_obs.add_run('Observaciones del horario: ')
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
p_hor_obs.add_run('_______________________________________________________________________________________________________________')


# ════════════════════════════════════════════════════════════════════════════
# ══ 8. COMPROMISOS ══════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '8.', 'COMPROMISOS Y OBSERVACIONES INICIALES')

for num, titulo, lineas in [
    ('8.1', 'Compromisos del Practicante (Política ITM Cap. VI, num. 6.3):', 3),
    ('8.2', 'Compromisos del Tutor (Escenario de práctica) – Art. 2.2.6.3.1.11 Decreto 0223/2026:', 3),
    ('8.3', 'Compromisos del Monitor ITM:', 3),
    ('8.4', 'Observaciones generales:', 2),
]:
    p_c = doc.add_paragraph()
    r = p_c.add_run(f'{num}  {titulo}')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
    for _ in range(lineas):
        p_l = doc.add_paragraph('_________________________________________________________________________________________________________')
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after  = Pt(1)


# ════════════════════════════════════════════════════════════════════════════
# ══ 9. REGISTRO FOTOGRÁFICO ═════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '9.', 'REGISTRO FOTOGRÁFICO DE LA REUNIÓN')

p_foto = doc.add_paragraph()
r1 = p_foto.add_run('Adjuntar foto de la reunión de inicio (opcional):  ')
r1.font.size = Pt(9); r1.font.color.rgb = NEGRO; r1.font.name = 'Arial'

p_auth = doc.add_paragraph()
r_a = p_auth.add_run(f'{CHK_OFF}  Autorizo el tratamiento de mis datos personales conforme a la Ley 1581/2012 y la Política de Privacidad del ITM.')
r_a.font.size = Pt(9); r_a.font.color.rgb = NEGRO; r_a.font.name = 'Arial'
p_pub = doc.add_paragraph()
r_p = p_pub.add_run(f'{CHK_OFF}  Autorizo la publicación de la fotografía en medios institucionales del ITM.')
r_p.font.size = Pt(9); r_p.font.color.rgb = NEGRO; r_p.font.name = 'Arial'

p_recuadro = doc.add_paragraph('[                     Espacio para pegar fotografía de la reunión                     ]')
p_recuadro.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_recuadro.runs[0].font.size = Pt(9); p_recuadro.runs[0].font.color.rgb = RGBColor(0x80,0x80,0x80)


# ════════════════════════════════════════════════════════════════════════════
# ══ 10. FECHA Y FIRMAS ══════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '10.', 'FECHA DE DILIGENCIAMIENTO Y FIRMAS')

p_fdig = doc.add_paragraph()
r = p_fdig.add_run('Fecha de diligenciamiento: ')
r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL; r.font.name = 'Arial'
p_fdig.add_run('_______________________')

tbl10 = doc.add_table(rows=6, cols=3)
tbl10.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl10, [6.0, 6.0, 6.0])

headers10 = ['PRACTICANTE', 'TUTOR DEL ESCENARIO', 'MONITOR ITM']
for i, h in enumerate(headers10):
    c = tbl10.rows[0].cells[i]
    set_cell_bg(c, AZUL)
    set_cell_borders(c, top='single', bottom='single', left='single', right='single', color='102D69')
    cell_para(c, h, bold=True, color=BLANCO, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

labels10 = [
    ('Nombre:', 'Nombre:', 'Nombre:'),
    ('Firma:', 'Firma:', 'Firma / Imagen:'),
    ('Documento:', 'Cargo:', ''),
    ('Programa:', 'Empresa:', ''),
    ('Correo:', 'Correo:', ''),
]
for ri, row_labels in enumerate(labels10):
    row = tbl10.rows[ri+1]
    for ci, lbl in enumerate(row_labels):
        c = row.cells[ci]
        alt = ri % 2 == 1
        set_cell_bg(c, GRIS_L if alt else BLANCO)
        set_cell_borders(c, top='single', bottom='single', left='single', right='single')
        if lbl:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            r_l = p.add_run(lbl + ' ')
            r_l.bold = True; r_l.font.size = Pt(8.5); r_l.font.color.rgb = AZUL; r_l.font.name = 'Arial'
            r_v = p.add_run('______________________________')
            r_v.font.size = Pt(8.5); r_v.font.color.rgb = NEGRO; r_v.font.name = 'Arial'
    set_row_height(row, 0.7)

set_row_height(tbl10.rows[2], 1.5)  # Firma, más alta


# ════════════════════════════════════════════════════════════════════════════
# ══ 11. MODALIDAD DE VISITA ═════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_section_bar(doc, '11.', 'MODALIDAD DE LA VISITA')

tbl11 = doc.add_table(rows=1, cols=2)
tbl11.alignment = WD_TABLE_ALIGNMENT.LEFT
set_col_widths(tbl11, [7.0, 11.0])
row11 = tbl11.rows[0]
set_cell_bg(row11.cells[0], AZUL_L)
set_cell_bg(row11.cells[1], BLANCO)
for c in row11.cells:
    set_cell_borders(c, top='single', bottom='single', left='single', right='single')
cell_para(row11.cells[0], 'Modalidad de la visita:', bold=True, color=AZUL, size=9)
p11 = row11.cells[1].paragraphs[0]
r = p11.add_run(f'{CHK_OFF} Presencial     {CHK_OFF} Virtual')
r.font.size = Pt(9); r.font.color.rgb = NEGRO; r.font.name = 'Arial'


# ── Referencia normativa pie ─────────────────────────────────────────────────
doc.add_paragraph()
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r_pie = p_pie.add_run(
    'Referencia normativa: Política de Prácticas Profesionales ITM · Decreto 0223 del 5 de marzo de 2026 '
    '(Min. Trabajo) · Ley 1780/2016 · Ley 2466/2025 (Reforma Laboral) · Ley 2365/2024 (Prevención acoso sexual) '
    '· Ley 2039/2020 (Experiencia profesional) · Art. 81 Código Sustantivo del Trabajo')
r_pie.italic = True; r_pie.font.size = Pt(7.5); r_pie.font.color.rgb = RGBColor(0x60,0x60,0x60); r_pie.font.name = 'Arial'


# ── Guardar ──────────────────────────────────────────────────────────────────
out_path = 'FDE_074_Guia_N1_Visita_Inicio_V11.docx'
doc.save(out_path)
print(f'Archivo generado: {out_path}')
